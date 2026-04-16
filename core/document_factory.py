import os
import zipfile
from reportlab.pdfgen import canvas
from docx import Document

class AssetFactory:
    """Generates the real digital products for the client."""
    
    def create_brand_pdf(self, company_name, strategy_text, output_path):
        """Generates an institutional-grade Brand Strategy PDF."""
        c = canvas.Canvas(output_path)
        c.setFont("Helvetica-Bold", 24)
        c.drawString(100, 750, f"BRAND STRATEGY: {company_name}")
        c.setFont("Helvetica", 12)
        c.drawString(100, 700, strategy_text)
        c.save()

    def create_brand_doc(self, company_name, niche, output_path):
        """Generates a professional Word Document for internal use."""
        doc = Document()
        doc.add_heading(f'{company_name} - Brand Identity', 0)
        doc.add_paragraph(f'Niche: {niche}')
        doc.add_paragraph('This document contains your automated brand voice guidelines.')
        doc.save(output_path)

    def package_assets(self, company_name, files_list, zip_name):
        """Zips everything into a clean deliverable package."""
        with zipfile.ZipFile(zip_name, 'w') as zipf:
            for file in files_list:
                zipf.write(file, os.path.basename(file))
        return zip_name
      
